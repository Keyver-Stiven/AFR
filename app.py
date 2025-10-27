import os
import json
import threading
import time
import traceback
from tkinter import messagebox
import customtkinter as ctk
from tkinter import filedialog
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from send2trash import send2trash

# ---------------- CONFIG ----------------
CONFIG_FILE = "config.json"

def cargar_config():
    defaults = {
        "ruta_formato": "", "carpeta_imagenes": "", "carpeta_destino": "",
        "grouping_option": "Por Fecha", "delete_option": "No Eliminar",
        "appearance_mode": "dark", "image_size": "Grande (3.0)"
    }
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                defaults.update(json.load(f))
        except Exception as e:
            print(f"Error al cargar config: {e}")
    return defaults

def guardar_config(cfg):
    try:
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(cfg, f, indent=4)
    except Exception as e:
        print(f"Error al guardar config: {e}")

config = cargar_config()
ctk.set_appearance_mode(config.get("appearance_mode", "dark"))
ctk.set_default_color_theme("dark-blue")

# ---------------- CLASE PARA EL MENÚ LATERAL ----------------
class SlidingPanel(ctk.CTkFrame):
    def __init__(self, parent, start_pos, end_pos):
        super().__init__(parent)
        self.start_pos, self.end_pos = start_pos, end_pos
        self.width = abs(start_pos - end_pos)
        self.pos = self.start_pos
        self.in_start_pos = True
        self.place(relx=self.start_pos, rely=0, relwidth=self.width, relheight=1)
        self.columnconfigure(0, weight=1)
        
        lbl_menu_title = ctk.CTkLabel(self, text="Opciones", font=ctk.CTkFont(size=16, weight="bold"))
        lbl_menu_title.grid(row=0, column=0, pady=(20, 10), padx=20, sticky="w")
        
        lbl_agrupar_por = ctk.CTkLabel(self, text="Agrupar imágenes por:", font=ctk.CTkFont(size=12))
        lbl_agrupar_por.grid(row=1, column=0, padx=20, pady=(15, 5), sticky="w")
        self.grouping_option = ctk.StringVar(value=config.get("grouping_option"))
        self.grouping_button = ctk.CTkSegmentedButton(self, values=["Por Fecha", "Por Nombre"], variable=self.grouping_option, command=lambda v: self.on_option_change("grouping_option", v))
        self.grouping_button.grid(row=2, column=0, padx=20, pady=5, sticky="ew")

        lbl_eliminar_fotos = ctk.CTkLabel(self, text="Al procesar las fotos:", font=ctk.CTkFont(size=12))
        lbl_eliminar_fotos.grid(row=3, column=0, padx=20, pady=(20, 5), sticky="w")
        self.delete_option = ctk.StringVar(value=config.get("delete_option"))
        radio_frame = ctk.CTkFrame(self, fg_color="transparent")
        radio_frame.grid(row=4, column=0, padx=20, pady=5, sticky="ew")
        radio_frame.columnconfigure(0, weight=1)
        rb1 = ctk.CTkRadioButton(radio_frame, text="No Eliminar", variable=self.delete_option, value="No Eliminar", command=lambda: self.on_option_change("delete_option", "No Eliminar"))
        rb2 = ctk.CTkRadioButton(radio_frame, text="Enviar a Papelera", variable=self.delete_option, value="A Papelera", command=lambda: self.on_option_change("delete_option", "A Papelera"))
        rb3 = ctk.CTkRadioButton(radio_frame, text="Eliminación Permanente", variable=self.delete_option, value="Permanente", command=lambda: self.on_option_change("delete_option", "Permanente"))
        rb1.grid(row=0, column=0, sticky="w", pady=4); rb2.grid(row=1, column=0, sticky="w", pady=4); rb3.grid(row=2, column=0, sticky="w", pady=4)

        lbl_img_size = ctk.CTkLabel(self, text="Tamaño de imagen (pulgadas):", font=ctk.CTkFont(size=12))
        lbl_img_size.grid(row=5, column=0, padx=20, pady=(20, 5), sticky="w")
        self.image_size_option = ctk.StringVar(value=config.get("image_size"))
        self.image_size_menu = ctk.CTkOptionMenu(self, variable=self.image_size_option, values=["Pequeño (2.0)", "Mediano (2.5)", "Grande (3.0)", "Muy Grande (3.5)"], command=lambda v: self.on_option_change("image_size", v))
        self.image_size_menu.grid(row=6, column=0, padx=20, pady=5, sticky="ew")

        lbl_tema = ctk.CTkLabel(self, text="Apariencia:", font=ctk.CTkFont(size=12))
        lbl_tema.grid(row=7, column=0, padx=20, pady=(20, 5), sticky="w")
        self.theme_switch = ctk.CTkSwitch(self, text="Tema Oscuro", command=self.cambiar_tema)
        self.theme_switch.grid(row=8, column=0, padx=20, pady=10, sticky="w")
        if config.get("appearance_mode") == "dark": self.theme_switch.select()

    def on_option_change(self, key, value): config[key] = value; guardar_config(config); print(f"Opción '{key}' cambiada a: {value}")
    def cambiar_tema(self): new_mode = "dark" if self.theme_switch.get() == 1 else "light"; ctk.set_appearance_mode(new_mode); self.on_option_change("appearance_mode", new_mode)
    def animate_forward(self):
        if self.pos < self.end_pos: self.pos += 0.02; self.place(relx=self.pos, rely=0, relwidth=self.width, relheight=1); self.after(10, self.animate_forward)
        else: self.in_start_pos = False
    def animate_backward(self):
        if self.pos > self.start_pos: self.pos -= 0.02; self.place(relx=self.pos, rely=0, relwidth=self.width, relheight=1); self.after(10, self.animate_backward)
        else: self.in_start_pos = True

# ---------------- UI & LAYOUT ----------------
root = ctk.CTk()
root.title("Automatización de Registro — AFR"); root.geometry("700x520"); root.resizable(False, False)
main_frame = ctk.CTkFrame(root, fg_color="transparent", width=560, height=520)
main_frame.place(relx=0.5, rely=0.5, anchor="center")

def animate_main_frame(target_relx):
    current_relx = float(main_frame.place_info().get('relx', "0.5"))
    if abs(target_relx - current_relx) > 0.01:
        new_relx = current_relx + (target_relx - current_relx) * 0.2
        main_frame.place(relx=new_relx, rely=0.5, anchor="center")
        root.after(10, lambda: animate_main_frame(target_relx))
    else: main_frame.place(relx=target_relx, rely=0.5, anchor="center")

card = ctk.CTkFrame(main_frame, corner_radius=14, width=520, height=480)
card.place(relx=0.5, rely=0.5, anchor="center")
slide_panel = SlidingPanel(root, start_pos=-0.4, end_pos=0)
def close_panel_if_open():
    if not slide_panel.in_start_pos: slide_panel.animate_backward(); animate_main_frame(0.5)
def handle_external_click(event): close_panel_if_open()
main_frame.bind("<Button-1>", handle_external_click); card.bind("<Button-1>", handle_external_click)
def toggle_slide_panel():
    if slide_panel.in_start_pos: slide_panel.animate_forward(); animate_main_frame(0.675)
    else: slide_panel.animate_backward(); animate_main_frame(0.5)

header_frame = ctk.CTkFrame(card, fg_color="transparent")
header_frame.pack(fill="x", padx=18, pady=(14,6))
btn_menu = ctk.CTkButton(header_frame, text="☰", width=30, command=toggle_slide_panel)
btn_menu.pack(side="left", padx=(0, 10))
lbl_title = ctk.CTkLabel(header_frame, text="⚙️ Automatización de Registro", font=ctk.CTkFont(size=16, weight="bold"))
lbl_title.pack(side="left")
header_frame.bind("<Button-1>", handle_external_click); lbl_title.bind("<Button-1>", handle_external_click)
sep = ctk.CTkFrame(card, height=1, fg_color=("gray80", "#222428"))
sep.pack(fill="x", padx=12, pady=(8,14))

def make_row(parent, label_text, var_initial=""):
    container = ctk.CTkFrame(parent, fg_color="transparent")
    container.pack(fill="x", padx=18, pady=6)
    lbl = ctk.CTkLabel(container, text=label_text, anchor="w", font=ctk.CTkFont(size=11))
    lbl.pack(fill="x", pady=(0,6))
    row = ctk.CTkFrame(container, fg_color="transparent")
    row.pack(fill="x")
    entry = ctk.CTkEntry(row, width=360, corner_radius=8)
    entry.insert(0, var_initial)
    entry.pack(side="left", padx=(0,8))
    btn = ctk.CTkButton(row, text="📁 Explorar", width=110)
    btn.pack(side="right")
    container.bind("<Button-1>", handle_external_click); lbl.bind("<Button-1>", handle_external_click); row.bind("<Button-1>", handle_external_click)
    return entry, btn

entry_formato, btn_formato = make_row(card, "Seleccionar Formato", config.get("ruta_formato"))
entry_img, btn_img = make_row(card, "Seleccionar carpeta de imágenes", config.get("carpeta_imagenes"))
entry_dest, btn_dest = make_row(card, "Seleccionar carpeta de destino", config.get("carpeta_destino"))

progress = ctk.CTkProgressBar(card, width=460)
lbl_status = ctk.CTkLabel(card, text="", font=ctk.CTkFont(size=12))
lbl_status.pack(pady=(6,0)); lbl_status.bind("<Button-1>", handle_external_click)
btn_iniciar = ctk.CTkButton(card, text="🚀 Iniciar proceso", width=420, corner_radius=12)
btn_iniciar.pack(pady=(18,12))
lbl_footer = ctk.CTkLabel(card, text="AFR — Automatización del Formato de Registro", font=ctk.CTkFont(size=10), text_color="gray")
lbl_footer.pack(side="bottom", pady=(0,12)); lbl_footer.bind("<Button-1>", handle_external_click)

def explorar_formato():
    ruta = filedialog.askopenfilename(title="Seleccionar formato (.docx)", filetypes=[("Word", "*.docx")])
    if ruta: entry_formato.delete(0, "end"); entry_formato.insert(0, ruta); config["ruta_formato"] = ruta; guardar_config(config)
def explorar_imagenes():
    ruta = filedialog.askdirectory(title="Seleccionar carpeta de imágenes")
    if ruta: entry_img.delete(0, "end"); entry_img.insert(0, ruta); config["carpeta_imagenes"] = ruta; guardar_config(config)
def explorar_destino():
    ruta = filedialog.askdirectory(title="Seleccionar carpeta de destino")
    if ruta: entry_dest.delete(0, "end"); entry_dest.insert(0, ruta); config["carpeta_destino"] = ruta; guardar_config(config)

btn_formato.configure(command=explorar_formato); btn_img.configure(command=explorar_imagenes); btn_dest.configure(command=explorar_destino)

# --- FUNCIÓN DE LIMPIEZA CORREGIDA ---
def clear_cell(cell):
    """Limpia todo el contenido de una celda, incluyendo imágenes y texto."""
    # Elimina las imágenes de la celda
    for p in cell.paragraphs:
        for run in p.runs:
            # Busca y elimina el elemento 'drawing' que contiene la imagen
            for drawing in run.element.xpath('.//w:drawing'):
                run.element.remove(drawing)
    # Elimina el texto de la celda
    for p in cell.paragraphs:
        p.clear()

def procesar_worker():
    try:
        ruta_formato = entry_formato.get().strip()
        carpeta_imagenes = entry_img.get().strip()
        carpeta_destino = entry_dest.get().strip()
        if not all([ruta_formato, carpeta_imagenes, carpeta_destino]):
            messagebox.showerror("Error", "Debes seleccionar todas las rutas."); lbl_status.configure(text="Faltan rutas", text_color="red"); return

        fotos = [os.path.join(carpeta_imagenes, f) for f in os.listdir(carpeta_imagenes) if f.lower().endswith((".jpg", ".jpeg", ".png"))]
        if not fotos:
            messagebox.showinfo("Info", "No se encontraron imágenes."); lbl_status.configure(text="Sin imágenes", text_color="orange"); return

        if slide_panel.grouping_option.get() == "Por Nombre": fotos.sort()
        else: fotos.sort(key=os.path.getctime)
        grupos = [fotos[i:i + 4] for i in range(0, len(fotos), 4)]
        
        progress.pack(pady=(18,6), before=lbl_status); progress.set(0)
        total_grupos = len(grupos)
        lbl_status.configure(text=f"Procesando {len(fotos)} imágenes...", text_color="orange")

        pat = re.compile(r"\((\d+)\)\.docx$")
        nums = [int(m.group(1)) for f in os.listdir(carpeta_destino) if f.lower().endswith(".docx") and (m := pat.search(f))]
        next_num = max(nums) + 1 if nums else 1

        size_str = slide_panel.image_size_option.get()
        try:
            img_size_val = float(re.search(r'\((\d+\.?\d*)\)', size_str).group(1))
        except (AttributeError, ValueError):
            img_size_val = 3.0

        for gi, grupo in enumerate(grupos):
            doc = Document(ruta_formato)
            tabla = doc.tables[0]
            posiciones = [(0,0), (0,1), (4,0), (4,1)]
            for r, c in posiciones:
                try: clear_cell(tabla.rows[r].cells[c])
                except Exception as e: print(f"Error al limpiar celda ({r},{c}): {e}")

            for idx, foto in enumerate(grupo):
                if idx >= len(posiciones): break
                try:
                    r, c = posiciones[idx]
                    par = tabla.rows[r].cells[c].paragraphs[0]
                    par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    par.add_run().add_picture(foto, width=Inches(img_size_val))
                except Exception as e: print(f"Error al insertar imagen {foto}: {e}")

            doc.save(os.path.join(carpeta_destino, f"{os.path.splitext(os.path.basename(ruta_formato))[0]} ({next_num}).docx"))

            delete_mode = slide_panel.delete_option.get()
            if delete_mode != "No Eliminar":
                for foto in grupo:
                    try:
                        # --- CORRECCIÓN DE RUTA DE ELIMINACIÓN ---
                        foto_normalizada = os.path.normpath(foto)
                        if delete_mode == "A Papelera":
                            send2trash(foto_normalizada)
                        elif delete_mode == "Permanente":
                            os.remove(foto_normalizada)
                    except Exception as e:
                        print(f"Error al eliminar foto {foto_normalizada}: {e}")
            
            next_num += 1
            progress.set((gi + 1) / total_grupos)
            lbl_status.configure(text=f"Procesando grupo {gi+1}/{total_grupos} ...")
            time.sleep(0.1)

        lbl_status.configure(text="Proceso finalizado ✅", text_color="green")
        messagebox.showinfo("Listo", "Proceso completado correctamente.")
        progress.pack_forget()
        config.update({"ruta_formato": ruta_formato, "carpeta_imagenes": carpeta_imagenes, "carpeta_destino": carpeta_destino})
        guardar_config(config)

    except Exception as e:
        traceback.print_exc()
        lbl_status.configure(text="Error durante el proceso", text_color="red")
        messagebox.showerror("Error", f"Ocurrió un error: {e}")
        progress.pack_forget()

def iniciar_proceso():
    t = threading.Thread(target=procesar_worker, daemon=True)
    t.start()

btn_iniciar.configure(command=iniciar_proceso)

root.mainloop()
